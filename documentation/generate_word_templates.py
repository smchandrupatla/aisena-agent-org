from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parent
OUT_ROOT = ROOT / "word-templates"

BRANDING_LINES = [
    "Brand: HSFS (Hybrid Sanctions and Fraud Screening System)",
    "Program: Autonomous Delivery Shop",
    "Status: Draft",
    "Classification: Internal",
    "Version: v1.0",
    "Last Updated: YYYY-MM-DD",
    "Owner: <Role/Name>",
    "Review Cycle: Quarterly",
]

TEMPLATES = {
    "01-planning-requirements": {
        "Business_Product_Requirements_Document_BRD_PRD": [
            "Executive Summary",
            "Problem Statement",
            "Business Objectives and KPIs",
            "Scope (In/Out)",
            "Personas and Stakeholders",
            "Functional Requirements",
            "Non-Functional Requirements",
            "Assumptions and Constraints",
            "Dependencies",
            "Acceptance Criteria",
            "Risks and Mitigations",
            "Sign-off",
        ],
        "Vision_and_Scope_Document": [
            "Vision Statement",
            "Strategic Goals",
            "Value Proposition",
            "Scope Boundaries",
            "Success Metrics",
            "Roadmap Summary",
        ],
        "User_Stories_Use_Cases": [
            "Story or Use Case ID",
            "Actor",
            "Narrative (As a / I want / So that)",
            "Preconditions",
            "Main Flow",
            "Alternate or Exception Flows",
            "Acceptance Criteria",
            "Traceability",
        ],
        "Functional_and_Non_Functional_Requirements": [
            "Functional Requirement Catalog",
            "Non-Functional Requirement Catalog",
            "Priority and MoSCoW Tag",
            "Verification Method",
            "Traceability Matrix",
        ],
        "Stakeholder_Analysis": [
            "Stakeholder Register",
            "Influence Interest Matrix",
            "Communication Plan",
            "Decision Rights and Escalation",
            "Risks from Stakeholder Misalignment",
        ],
    },
    "02-design": {
        "System_Architecture_Document": [
            "Context and Drivers",
            "Architecture Principles",
            "Logical Architecture",
            "Physical Deployment Architecture",
            "Data Flow and Integration Points",
            "Security Reliability and Performance",
            "Trade-offs and Open Questions",
        ],
        "High_Level_Design_HLD": [
            "Service Boundaries",
            "Major Components",
            "External Interfaces",
            "Key Data Contracts",
            "Scalability Overview",
        ],
        "Low_Level_Design_LLD": [
            "Module and Class Breakdown",
            "Sequence Design",
            "Error Handling",
            "Validation Rules",
            "Logging and Observability",
        ],
        "Database_Schema_ER_Diagrams": [
            "ER Diagram Overview",
            "Entity Definitions",
            "Keys Constraints and Indexes",
            "Retention and Lifecycle",
            "Migration Strategy",
        ],
        "API_Specifications_OpenAPI_Swagger": [
            "API Purpose and Scope",
            "Endpoint Catalog",
            "Request Response Schemas",
            "Authentication and Authorization",
            "Error Model",
            "Versioning and Deprecation",
        ],
        "UI_UX_Wireframes_and_Design_System": [
            "User Journey Map",
            "Wireframes by Screen",
            "Design System Components",
            "Accessibility Requirements",
            "Interaction States",
        ],
        "Architecture_Decision_Records_ADR": [
            "Decision Title",
            "Status",
            "Context",
            "Decision",
            "Alternatives Considered",
            "Consequences",
            "Follow-up Actions",
        ],
    },
    "03-development": {
        "Coding_Standards_Style_Guide": [
            "Language-Specific Conventions",
            "Naming and Structure Rules",
            "Error Handling Standards",
            "Logging Standards",
            "Code Review Checklist",
        ],
        "Code_Comments_and_Inline_Documentation": [
            "Commenting Principles",
            "Function and Class Documentation Format",
            "Public API Docstring Requirements",
            "Examples and Anti-Patterns",
        ],
        "README_Per_Repo_Module": [
            "Purpose",
            "Architecture Summary",
            "Quick Start",
            "Configuration",
            "Build and Test Commands",
            "Troubleshooting",
        ],
        "Setup_and_Environment_Configuration_Guide": [
            "Prerequisites",
            "Environment Variables",
            "Local Setup Steps",
            "Verification Checks",
            "Common Issues and Fixes",
        ],
        "Third_Party_Library_Dependency_Documentation": [
            "Dependency Inventory",
            "Version and License",
            "Security Considerations",
            "Upgrade Policy",
            "Replacement or Fallback Strategy",
        ],
    },
    "04-testing-qa": {
        "Test_Plan": [
            "Scope and Objectives",
            "Test Strategy and Levels",
            "Entry Exit Criteria",
            "Environment and Test Data",
            "Schedule and Ownership",
            "Risks and Mitigation",
        ],
        "Test_Cases_and_Test_Scripts": [
            "Test Case ID",
            "Preconditions",
            "Steps",
            "Expected Results",
            "Actual Results",
            "Status and Evidence",
        ],
        "Bug_Defect_Reports": [
            "Defect ID and Title",
            "Severity Priority",
            "Environment",
            "Reproduction Steps",
            "Expected vs Actual",
            "Root Cause and Fix Reference",
        ],
        "Test_Coverage_Reports": [
            "Coverage Scope",
            "Requirement Traceability",
            "Unit Integration E2E Coverage",
            "Gaps and Risk Impact",
            "Improvement Plan",
        ],
        "UAT_Sign_Off_Documents": [
            "UAT Scope",
            "Participant List",
            "Executed Scenarios",
            "Open Issues",
            "Formal Sign-off Decision",
        ],
    },
    "05-devops-infrastructure": {
        "CI_CD_Pipeline_Documentation": [
            "Pipeline Overview",
            "Stages and Gates",
            "Quality and Security Checks",
            "Artifact Promotion Flow",
            "Failure Handling",
        ],
        "Deployment_Guides_Runbooks": [
            "Pre-deployment Checks",
            "Deployment Steps",
            "Verification Steps",
            "Rollback Steps",
            "Escalation Contacts",
        ],
        "Infrastructure_as_Code_Documentation": [
            "IaC Scope",
            "Module Layout",
            "Variables and State Strategy",
            "Security Controls",
            "Change and Review Process",
        ],
        "Environment_Configuration_Dev_Staging_Prod": [
            "Environment Matrix",
            "Configuration Differences",
            "Secret Handling Strategy",
            "Drift Detection",
            "Approval Workflow",
        ],
        "Monitoring_and_Alerting_Setup": [
            "SLOs and SLIs",
            "Metrics and Logs",
            "Alert Rules and Thresholds",
            "On-call Routing",
            "Dashboard References",
        ],
        "Disaster_Recovery_Rollback_Plans": [
            "Recovery Objectives (RTO/RPO)",
            "Failure Scenarios",
            "Recovery Procedures",
            "Backup Restore Validation",
            "Drill Schedule",
        ],
    },
    "06-security-compliance": {
        "Security_Architecture_Threat_Model": [
            "System Context and Assets",
            "Threat Identification",
            "Attack Surface Analysis",
            "Mitigation Controls",
            "Residual Risk Acceptance",
        ],
        "Data_Privacy_and_Compliance_Documentation": [
            "Applicable Regulations",
            "Data Classification and Flows",
            "Processing Purpose and Lawful Basis",
            "Retention and Deletion Policy",
            "Compliance Evidence and Audits",
        ],
        "Access_Control_and_Permissions_Matrix": [
            "Role Inventory",
            "Privilege Matrix",
            "Segregation of Duties",
            "Approval and Provisioning Workflow",
            "Review and Recertification Cadence",
        ],
        "Audit_Logs_Policy": [
            "Events to Log",
            "Log Integrity and Tamper Controls",
            "Retention and Storage",
            "Access and Review Process",
            "Incident Investigation Workflow",
        ],
    },
    "07-project-management": {
        "Project_Charter": [
            "Purpose and Objectives",
            "Scope and Deliverables",
            "Timeline and Milestones",
            "Roles and Responsibilities",
            "Governance and Approvals",
        ],
        "Sprint_Iteration_Plans_and_Backlogs": [
            "Sprint Goal",
            "Committed Stories Tasks",
            "Capacity and Allocation",
            "Dependencies and Risks",
            "Definition of Done",
        ],
        "Meeting_Notes_and_Decision_Logs": [
            "Meeting Metadata",
            "Agenda and Discussion Summary",
            "Decisions and Owners",
            "Action Items",
            "Follow-up Date",
        ],
        "Risk_Register": [
            "Risk ID and Description",
            "Impact and Probability",
            "Mitigation Strategy",
            "Owner and Due Date",
            "Current Status",
        ],
        "Status_Reports_Stakeholder_Updates": [
            "Executive Summary",
            "Progress Against Plan",
            "KPI Snapshot",
            "Blockers and Escalations",
            "Next Period Plan",
        ],
    },
    "08-release-maintenance": {
        "Release_Notes_Changelog": [
            "Release Identifier and Date",
            "Features Added or Changed",
            "Fixes and Improvements",
            "Breaking Changes",
            "Known Limitations",
        ],
        "Versioning_Documentation": [
            "Versioning Scheme",
            "Branching and Tagging Rules",
            "Compatibility Policy",
            "Upgrade Path",
            "Deprecation Window",
        ],
        "Known_Issues_Log": [
            "Issue ID and Description",
            "Impact and Workaround",
            "Affected Versions",
            "Owner and ETA",
            "Resolution Status",
        ],
        "Post_Release_Support_Maintenance_Guide": [
            "Support Scope",
            "SLA Response Targets",
            "Triage Process",
            "Hotfix Procedure",
            "Escalation Matrix",
        ],
        "End_of_Life_Deprecation_Notices": [
            "Deprecated Component",
            "Timeline and Milestones",
            "Migration Guidance",
            "Communication Plan",
            "Final Retirement Steps",
        ],
    },
    "09-end-user": {
        "User_Manual_Help_Documentation": [
            "Audience and Scope",
            "Feature Walkthroughs",
            "Step-by-step Procedures",
            "Error Messages and Recovery",
            "Support Contacts",
        ],
        "FAQ": [
            "General Questions",
            "Account and Access",
            "Troubleshooting",
            "Security and Privacy",
            "Escalation Paths",
        ],
        "Training_Materials": [
            "Learning Objectives",
            "Training Agenda",
            "Practical Exercises",
            "Assessment Criteria",
            "Reference Material",
        ],
        "API_Developer_Documentation": [
            "API Overview",
            "Authentication",
            "Endpoint Quick Start",
            "Request Response Examples",
            "Rate Limits and Error Codes",
        ],
    },
    "10-handover-knowledge-transfer": {
        "System_Handover_Document": [
            "System Summary",
            "Architecture and Components",
            "Operational Procedures",
            "Open Risks and Known Issues",
            "Ownership Transition Checklist",
        ],
        "Onboarding_Guide_for_New_Developers": [
            "Access and Tooling Setup",
            "Repository Structure",
            "Build Test Run Commands",
            "Contribution Workflow",
            "Common Pitfalls",
        ],
        "Glossary_of_Terms_Domain_Knowledge": [
            "Domain Terms",
            "Acronyms",
            "Business Rules Vocabulary",
            "Data and Event Definitions",
            "Source References",
        ],
    },
}


def add_branding(doc: Document, doc_type: str) -> None:
    doc.add_heading("HSFS (Hybrid Sanctions and Fraud Screening System)", level=1)
    doc.add_paragraph("Autonomous Delivery Shop")
    doc.add_paragraph(f"Document Type: {doc_type}")
    for line in BRANDING_LINES:
        doc.add_paragraph(line)
    doc.add_paragraph("-" * 60)


def add_sections(doc: Document, sections: list[str]) -> None:
    for idx, title in enumerate(sections, start=1):
        doc.add_heading(f"{idx}. {title}", level=2)
        doc.add_paragraph("[Add content here]")


def add_footer(doc: Document) -> None:
    doc.add_paragraph("HSFS | Controlled Document | Do not distribute without approval")


def create_template(path: Path, doc_type: str, sections: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    add_branding(doc, doc_type)
    add_sections(doc, sections)
    add_footer(doc)
    doc.save(path)


def main() -> None:
    total = 0
    for category, templates in TEMPLATES.items():
        category_dir = OUT_ROOT / category
        for name, sections in templates.items():
            doc_type = name.replace("_", " ")
            target = category_dir / f"TEMPLATE-{name}.docx"
            create_template(target, doc_type, sections)
            total += 1
    print(f"Generated {total} Word templates under {OUT_ROOT}")


if __name__ == "__main__":
    main()
